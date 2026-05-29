#!/usr/bin/env python3
"""Genererer et Word-dokument (liggende) med all brukervendt tekst fra straverso.com.

Tabell: ID | Seksjon | Element | Dagens tekst | Ny tekst
Fyll ut «Ny tekst» – du kan bruke punktlister og nummererte lister der.
La en celle stå tom for å beholde dagens tekst.
"""

import os
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from content_data import ROWS

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

INDIGO = "0F1354"
OFFWHITE = "F5F2EB"
NEWCOL = "FFF4F4"  # svak coral for «Ny tekst»-kolonnen

# ID | Seksjon | Element | Dagens tekst | Ny tekst  (bredder i cm, sum ~ 26 cm)
COL_WIDTHS = [Cm(3.0), Cm(2.6), Cm(4.6), Cm(8.0), Cm(8.0)]
HEADERS = ["ID", "Seksjon", "Element", "Dagens tekst", "Ny tekst"]


def shade_cell(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=None, size=None, italic=False):
    cell.text = ""
    para = cell.paragraphs[0]
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            para.add_run().add_break()
        run = para.add_run(line)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if size:
            run.font.size = Pt(size)


def set_col_widths(table):
    # Sett bredde på hver celle (Word respekterer dette best når autofit er av)
    for row in table.rows:
        for idx, width in enumerate(COL_WIDTHS):
            row.cells[idx].width = width


doc = Document()

# Liggende A4 med smale marger
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# Tittel + kort instruksjon
h = doc.add_heading("Straverso.com – innhold", level=1)
intro = doc.add_paragraph()
intro.add_run(
    "Fyll inn ønsket tekst i kolonnen «Ny tekst». Du kan bruke punktlister og "
    "nummererte lister der det passer. La en celle stå tom for å beholde dagens "
    "tekst. Ikke endre ID-kolonnen – den brukes til å plassere teksten riktig på siden."
).italic = True

table = doc.add_table(rows=1, cols=len(HEADERS))
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
table.allow_autofit = False

# Header-rad
hdr = table.rows[0].cells
for i, text in enumerate(HEADERS):
    set_cell_text(hdr[i], text, bold=True, color=OFFWHITE)
    shade_cell(hdr[i], INDIGO)

# Datarader
for rid, sec, element, current in ROWS:
    cells = table.add_row().cells
    set_cell_text(cells[0], rid, size=8, color="888888")
    set_cell_text(cells[1], sec)
    set_cell_text(cells[2], element)
    set_cell_text(cells[3], current)
    set_cell_text(cells[4], "")  # Ny tekst – fylles av bruker
    shade_cell(cells[4], NEWCOL)

set_col_widths(table)

# Gjenta header-raden øverst på hver side
tr = table.rows[0]._tr
trpr = tr.get_or_add_trPr()
th = OxmlElement("w:tblHeader")
th.set(qn("w:val"), "true")
trpr.append(th)

out = os.path.join(ROOT, "straverso-innhold.docx")
doc.save(out)
print(f"Skrev {out} med {len(ROWS)} tekstrader.")
