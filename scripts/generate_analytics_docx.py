#!/usr/bin/env python3
"""Genererer et Word-dokument som forklarer analytics-oppsettet for straverso.com.

Dekker: Vercel Web Analytics, Vercel Speed Insights, og custom events satt opp i koden.
Brukes som internt referansedokument.
"""

import os
import sys
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

INDIGO = "0F1354"
OFFWHITE = "F5F2EB"
ACCENT = "FF6B6B"
MUTED = "F5F5F8"


def shade_cell(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def set_cell(cell, text, *, bold=False, color=None, size=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        if size:
            r.font.size = Pt(size)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, *, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows, widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, color=OFFWHITE)
        shade_cell(hdr[i], INDIGO)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val)
    if widths_cm:
        for row in table.rows:
            for i, w in enumerate(widths_cm):
                row.cells[i].width = Cm(w)
    return table


# ─── Build document ─────────────────────────────────────────────
doc = Document()

# Page setup – A4
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

# Title
title = doc.add_heading("Straverso.com – Analytics og bruksdata", level=0)
sub = doc.add_paragraph()
sub_run = sub.add_run("Hvordan vi måler bruk av nettsiden, hva tallene betyr, og hvordan du kan bruke dem.")
sub_run.italic = True
sub_run.font.color.rgb = RGBColor.from_string("555555")

# 1. Oversikt
add_heading(doc, "1. Oversikt over hva som er installert", level=1)
add_para(doc,
    "Du har tre uavhengige målesystemer aktive på straverso.com. Alle "
    "kjører kun i produksjon (ikke under lokal utvikling), er gratis i Vercel "
    "sitt gratisnivå, og krever ingen samtykkebanner siden de ikke bruker "
    "cookies eller sporer enkeltbrukere."
)
add_bullets(doc, [
    "Vercel Web Analytics – overordnet trafikk og publikum (passiv, anonym).",
    "Vercel Speed Insights – ytelse målt på ekte besøkere (Core Web Vitals).",
    "Custom events – hendelser i koden som forteller hva folk klikker på.",
])

# 2. Vercel Web Analytics
add_heading(doc, "2. Vercel Web Analytics", level=1)
add_para(doc,
    "Du finner dette under Project → Analytics-fanen i Vercel-dashboardet "
    "etter første deploy. Data oppdateres med ~1 minutts forsinkelse."
)

add_heading(doc, "Spørsmål du får svar på", level=2)
add_table(doc,
    headers=["Måling", "Eksempel-innsikt"],
    rows=[
        ("Sidevisninger og besøkende", "143 besøk i går, +22 % vs. forrige uke."),
        ("Top sider", "/ (eneste i dag, men anchor-navigasjon kan trackes senere)."),
        ("Hvor trafikken kommer fra", "LinkedIn 34 %, Google 28 %, direkte 22 %, andre 16 %."),
        ("Top referrers (URL-nivå)", "Hvilken Slack-tråd eller artikkel som sender folk."),
        ("Geografi", "67 % Norge, 14 % USA, 9 % Sverige, …"),
        ("Enheter og nettlesere", "Desktop 58 %, mobil 38 %, tablet 4 %. Safari 41 %, Chrome 38 %, …"),
        ("Operativsystem", "macOS 35 %, iOS 28 %, Windows 24 %, Android 10 %."),
    ],
    widths_cm=[5.5, 11.0],
)

# 3. Speed Insights
add_heading(doc, "3. Vercel Speed Insights", level=1)
add_para(doc,
    "Egen fane i Vercel under Project → Speed Insights. Måler de samme Web "
    "Vitals som Google bruker til SEO-rangering, så det dobler som ytelse- "
    "og SEO-tilbakemelding."
)

add_heading(doc, "Hva hver måling betyr", level=2)
add_table(doc,
    headers=["Måling", "Hva det betyr", "Mål"],
    rows=[
        ("LCP", "Largest Contentful Paint – hvor raskt hovedinnholdet vises.", "< 2.5 s"),
        ("CLS", "Cumulative Layout Shift – om elementer hopper rundt mens siden laster.", "< 0.1"),
        ("INP", "Interaction to Next Paint – hvor raskt UI responderer på klikk.", "< 200 ms"),
        ("FCP", "First Contentful Paint – når noe synlig vises i det hele tatt.", "< 1.8 s"),
        ("TTFB", "Time to First Byte – serverrespons.", "< 600 ms"),
    ],
    widths_cm=[2.5, 11.0, 3.0],
)

add_para(doc,
    "Dashboardet viser median og 75. persentil per device-type (mobil / "
    "desktop). Etter hver deploy får du regresjons-deteksjon: «LCP økte "
    "med 400 ms etter forrige deploy»."
)

# 4. Custom events
add_heading(doc, "4. Custom events i koden", level=1)
add_para(doc,
    "Fire egendefinerte hendelser er instrumentert. Hver gang en bruker "
    "trigger en av dem, sendes en hendelse til Vercel Analytics og dukker "
    "opp under Events-fanen med eventuell metadata."
)

add_table(doc,
    headers=["Event-navn", "Når den triggres", "Metadata", "Hva den forteller deg"],
    rows=[
        ("hero_cta",
         "Klikk på «Utforsk produktene» eller «Ta kontakt» i hero.",
         "which: products | contact\nlang: no | en",
         "Hvor stor andel scroller selv ned vs. trykker CTA. Splitter NO/EN."),
        ("contact_click",
         "Klikk på en av de to mailto-knappene i kontakt-seksjonen.",
         "variant: cta | address\nlang: no | en",
         "Konverteringsrate. «cta» = primærknapp, «address» = epost-link."),
        ("lang_toggle",
         "Brukeren bytter mellom NO og EN i nav-en.",
         "from: no | en\nto: no | en",
         "Hvor stor andel av besøk er engelskspråklig. Mister du folk i toggle?"),
        ("products_viewed",
         "Produkter-seksjonen kommer i viewport (scroll-dybde).",
         "lang: no | en",
         "Andel besøk som scroller forbi hero. Indikerer interesse."),
    ],
    widths_cm=[3.5, 4.5, 3.5, 5.5],
)

# 5. Hvor du finner dataene
add_heading(doc, "5. Hvor du finner dataene", level=1)
add_bullets(doc, [
    "vercel.com → Project «straverso.com» → Analytics-fanen (sidevisninger, kilder, geo, devices).",
    "Samme prosjekt → Speed Insights-fanen (Core Web Vitals over tid).",
    "Analytics-fanen → seksjonen «Custom Events» nederst (de fire event-typene ovenfor).",
    "Alle data lar seg filtrere på tidsrom (24t / 7d / 30d / 90d) og på land, device, source.",
])

# 6. Praktisk bruk – eksempler
add_heading(doc, "6. Praktisk bruk – eksempel-spørsmål du kan svare på", level=1)

add_heading(doc, "Markedsføring", level=2)
add_bullets(doc, [
    "«Ga LinkedIn-posten min i går faktisk trafikk?» → Analytics → Sources → filtrer 24t.",
    "«Hvilken kanal konverterer best til faktisk kontakt?» → Sammenlign besøkere per kilde mot contact_click-events.",
    "«Når besøker folk siden mest – hvilken ukedag?» → Analytics → tidsrom-graf.",
])

add_heading(doc, "Produkt", level=2)
add_bullets(doc, [
    "«Hvor stor andel forstår at det er nordmenn de besøker?» → Andel lang_toggle til en av total besøk fra utlandet.",
    "«Scroller folk forbi hero?» → products_viewed / total visits.",
    "«Hvilken hero-CTA virker?» → Sammenlign hero_cta (products) vs (contact).",
])

add_heading(doc, "Ytelse og SEO", level=2)
add_bullets(doc, [
    "«Er siden rask nok for Google?» → Speed Insights → alle Web Vitals i grønt.",
    "«Hva er flaskehalsen på mobil?» → Filter på Mobile → se hvilken metric som er rødest.",
    "«Knekker forrige deploy noe?» → Speed Insights viser endring etter hver deploy.",
])

# 7. Tips for å utvide
add_heading(doc, "7. Tips for å utvide senere", level=1)
add_bullets(doc, [
    "Flere custom events: track('newsletter_signup', …), track('product_card_click', { name }), track('scroll_depth_75'), …",
    "Goals / funnels: krever GA4 eller PostHog (gratis nivå er rikelig for trafikkvolumet ditt).",
    "A/B-test av tagline eller CTA-tekst: kan settes opp med Vercel Edge Config + et lite custom event for å måle vinner.",
    "Heat-maps og session replay: Microsoft Clarity er gratis og enkelt å koble på.",
])

# 8. Personvern / GDPR
add_heading(doc, "8. Personvern og GDPR", level=1)
add_para(doc,
    "Vercel Web Analytics + Speed Insights samler ingen personidentifiserbare "
    "data, bruker ingen cookies, og krever ingen samtykkebanner. IP-adresser "
    "blir hash-et og anonymisert før behandling. Custom events inneholder "
    "kun de feltene du selv legger inn (lang, which, source) – aldri navn "
    "eller e-post."
)

# Footer note
doc.add_paragraph()
note = doc.add_paragraph()
note_run = note.add_run("Dokument generert automatisk. Kjør scripts/generate_analytics_docx.py for å oppdatere.")
note_run.italic = True
note_run.font.size = Pt(9)
note_run.font.color.rgb = RGBColor.from_string("888888")

# Save
out = os.path.join(ROOT, "straverso-analytics.docx")
doc.save(out)
print(f"Skrev {out}")
